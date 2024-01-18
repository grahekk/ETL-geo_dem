from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

europa_image_path = '/home/nikola/4_north_america/GeoDataPump/scripts/tree_cover_density.vrt_graphics.jpeg'  
hrvatska_image_path = '/home/nikola/4_north_america/GeoDataPump/scripts/tree_cover_density.vrt_graphics.jpeg' 

def create_dem_documentation():
    # Create a new Document
    doc = Document()

    # Add Title
    title = "NA DEM Documentation"
    doc.add_heading(title, level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add Heading
    heading = "1. SLOPE_100m_EUROPE"
    doc.add_heading(heading, level=1)

    # Add Table
    table_data = {
        "Verzija": "1.0",
        "Datum promjene": "13.2.2023.",
        "Kratki opis": "Raster izveden iz rastera digital surface model rezolucije 100m, vrijednosti pixela su stupnjevi nagiba terena u odnosu na horizontalu površinu na koju je gravitacijska sila okomita.",
        "Rezolucija (horizontalna)": "100 m",
        "Referentni koordinatni sustav": "EPSG:3035 - ETRS89-extended / LAEA Europe",
        "Parametri": "ratio vertical to horizontal units = 1; mjerna jedinica = stupanj [°]; No data = -9999 (u domeni), None (van domene)",
        "Vrijeme izrade konačnog file-a": "~ 1 minuta",
        "Path do file-a sa podacima": "/mnt/volume-nbg1-1/satellite/eu_slope",
        "Path do skripte": "/home/domagoj/OPENSTREETMAP/FINAL_OSM.py",
        "Funkcija za izradu file-a": "slope_aspect(lat,lon,slope = True, aspect = False, elevation = False, Europe = True); koristi se gdal:slope",
        "Slike dataseta (Europa)": f"(image:{europa_image_path})",
        "Slike dataseta (Hrvatska)": f"(image:{hrvatska_image_path})",
        "Web download data": "/",
        "Metadata": "/mnt/volume-nbg1-1/satellite/eu_slope/slope_metadata.txt",
    }

    table = doc.add_table(rows=len(table_data), cols=2)
    for row, (key, value) in enumerate(table_data.items()):
        table.cell(row, 0).text = key
        if "(image:" in value and value.endswith(")"):
            start_index = value.find("(image:") + len("(image:")
            end_index = value.find(")")
            image_path = value[start_index:end_index].strip()
            cell = table.cell(row, 1)
            paragraph = cell.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(2))
        else:
            table.cell(row, 1).text = value



    # Save the document
    doc.save('dem_documentation.docx')

if __name__ == "__main__":
    create_dem_documentation()
    print("Doc done")