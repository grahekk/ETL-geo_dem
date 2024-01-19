import os
import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

metadata_list_path = "/home/nikola/4_north_america/GeoDataPump/scripts/docs/metadata_list.json"
with open(metadata_list_path, "r") as file:
    metadata_list = json.load(file)

def create_dem_documentation():
    # Create a new Document
    doc = Document()

    # Add Title
    title = "NA DEM Documentation"
    doc.add_heading(title, level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Iterate over metadata entries
    for metadata_entry in metadata_list:
        heading = os.path.basename(metadata_entry["file_path"])
        doc.add_page_break()
        doc.add_heading(heading, level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add Table
        table_data = {
            "Verzija": metadata_entry["version"],
            "Datum promjene": f"Date created = {metadata_entry['date_created']}, Date_modified = {metadata_entry['date_modified']}",
            "Veličina": metadata_entry["file_size"],
            "Kratki opis": metadata_entry["description"],
            "Rezolucija (horizontalna)": metadata_entry["resolution"],
            "Referentni koordinatni sustav": metadata_entry["coordinate_system"],
            "Atributi": metadata_entry["attributes"],
            "Vrijeme izrade konačnog file-a": metadata_entry["time_to_make"],
            "Path do file-a sa podacima": metadata_entry["file_path"],
            "Path do skripte": metadata_entry["script_path"],
            "Funkcija za izradu file-a": metadata_entry["creation_method"],
            "Web download data": metadata_entry["download_url"],
            f"Slike dataseta ({heading})": f'(image:{metadata_entry["image_path"]})',
        }

        table = doc.add_table(rows=len(table_data), cols=2)
        for row, (key, value) in enumerate(table_data.items()):
            table.cell(row, 0).text = key
            if "(image:" in value and value.endswith(")"):
                start_index = value.find("(image:") + len("(image:")
                end_index = value.find(")")
                image_path = value[start_index:end_index].strip()
                cell = table.cell(row, 1)
                paragraph = cell.add_paragraph()
                run = paragraph.add_run()
                try:
                    run.add_picture(image_path, width=Inches(4.5))
                except FileNotFoundError:
                    table.cell(row, 1).text = "no picture"
            else:
                cell = table.cell(row, 1)
                cell.text = value

                # Set alignment to justify
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Save the document
    doc.save('/home/nikola/4_north_america/GeoDataPump/scripts/docs/dem_documentation.docx')

if __name__ == "__main__":
    create_dem_documentation()
    print("Doc done")