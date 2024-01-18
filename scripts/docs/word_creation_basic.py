from docx import Document
from docx.shared import Inches


img_path = "/home/nikola/4_north_america/GeoDataPump/scripts/tree_cover_density.vrt_graphics.jpeg"

def create_docx_with_table_and_image():
    # Create a new Document
    doc = Document()

    # Add some text
    doc.add_heading('Document with Table and Image', level=1)
    doc.add_paragraph('This is some text in the document.')

    # Add a table
    table = doc.add_table(rows=3, cols=3)
    for i in range(3):
        for j in range(3):
            table.cell(i, j).text = f'Row {i+1}, Col {j+1}'

    # Add an image to the first cell of the table
    table.cell(0, 0).paragraphs[0].add_run().add_picture(img_path, width=Inches(1.5))

    # Add Images
    doc.add_paragraph("Europa:")
    doc.add_picture(img_path, width=Inches(5))

    doc.add_paragraph("Hrvatska:")
    doc.add_picture(img_path, width=Inches(5))

    # Save the document
    doc.save('document_with_table_and_image.docx')

if __name__ == "__main__":
    create_docx_with_table_and_image()
    print
